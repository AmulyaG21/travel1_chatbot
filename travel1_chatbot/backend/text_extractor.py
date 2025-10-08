from docx import Document
import pdfplumber
from pathlib import Path
from typing import List, Dict, Optional


def extract_text_from_file(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()

    elif ext == ".docx":
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text.strip()

    elif ext == ".pdf":
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts).strip()

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_segments(file_path: str) -> List[Dict[str, Optional[str]]]:
    """Return a list of segments with source filename and page numbers for citation.
    Each segment dict contains: {"text": str, "page": Optional[int], "source": str}
    """
    src_name = Path(file_path).name
    ext = Path(file_path).suffix.lower()
    segments: List[Dict[str, Optional[str]]] = []

    if ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    segments.append({
                        "text": page_text.strip(),
                        "page": idx,
                        "source": src_name,
                    })
    elif ext == ".docx":
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs]).strip()
        if text:
            segments.append({
                "text": text,
                "page": None,  # DOCX has no simple page mapping here
                "source": src_name,
            })
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
            if text:
                segments.append({
                    "text": text,
                    "page": None,
                    "source": src_name,
                })
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return segments
